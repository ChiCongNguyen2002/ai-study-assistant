from .base import DocumentParser
from docx import Document

class DocxParser(DocumentParser):
    def parse(self, file_path: str) -> dict:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        return {
            'title': file_path.split('/')[-1],
            'content': text,
            'pages': len(doc.paragraphs),
            'metadata': {'format': 'docx', 'type': 'document'}
        }
