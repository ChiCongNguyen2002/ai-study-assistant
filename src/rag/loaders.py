"""Document loaders for RAG system."""
import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document as DocxDocument
from datetime import datetime


class PDFLoader:
    """Load and parse PDF documents."""

    @staticmethod
    def load_pdf(file_path: str) -> List[Dict]:
        """
        Load PDF and extract text per page, with metadata.

        Note: pages are returned separately (not pre-chunked) so the
        chunker can merge them per source document and split on
        paragraph boundaries instead of hard page cuts.

        Args:
            file_path: Path to PDF file

        Returns:
            List of per-page entries with metadata
        """
        pages = []
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                filename = Path(file_path).name

                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        pages.append({
                            "content": text,
                            "source": filename,
                            "page": page_num,
                            "type": "pdf",
                            "loaded_at": datetime.now().isoformat(),
                            "sensitivity": "MEDIUM"  # Default, will be updated by security scan
                        })
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")

        return pages

    @staticmethod
    def load_all_pdfs(pdf_directory: str) -> List[Dict]:
        """Load all PDFs from directory."""
        all_pages = []
        pdf_dir = Path(pdf_directory)

        if not pdf_dir.exists():
            print(f"PDF directory not found: {pdf_directory}")
            return all_pages

        for pdf_file in pdf_dir.glob("**/*.pdf"):
            print(f"Loading: {pdf_file.name}")
            pages = PDFLoader.load_pdf(str(pdf_file))
            all_pages.extend(pages)

        print(f"✓ Loaded {len(all_pages)} pages from PDFs")
        return all_pages


class DocxLoader:
    """Load and parse Word (.docx) documents."""

    @staticmethod
    def load_docx(file_path: str) -> List[Dict]:
        """
        Load a .docx file and extract text.

        Word documents have no native page concept, so the whole
        document is returned as a single entry (page=0); the chunker
        splits it on paragraph boundaries same as merged PDF text.

        Args:
            file_path: Path to .docx file

        Returns:
            List with a single document entry (empty if load failed)
        """
        entries = []
        try:
            filename = Path(file_path).name
            doc = DocxDocument(file_path)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        paragraphs.append(row_text)

            content = "\n\n".join(paragraphs)
            if content.strip():
                entries.append({
                    "content": content,
                    "source": filename,
                    "page": 0,
                    "type": "docx",
                    "loaded_at": datetime.now().isoformat(),
                    "sensitivity": "MEDIUM"
                })
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")

        return entries

    @staticmethod
    def load_all_docx(directory: str) -> List[Dict]:
        """Load all .docx files from directory."""
        all_entries = []
        doc_dir = Path(directory)

        if not doc_dir.exists():
            print(f"Document directory not found: {directory}")
            return all_entries

        for docx_file in doc_dir.glob("**/*.docx"):
            print(f"Loading: {docx_file.name}")
            all_entries.extend(DocxLoader.load_docx(str(docx_file)))

        print(f"✓ Loaded {len(all_entries)} Word documents")
        return all_entries


class ConfluenceLoader:
    """Load documents from Confluence wiki."""

    def __init__(self, base_url: str, username: str, api_token: str):
        """Initialize Confluence loader."""
        self.base_url = base_url
        self.username = username
        self.api_token = api_token

    def load_confluence_pages(self, space_key: str = "TECH") -> List[Dict]:
        """
        Load pages from Confluence space.

        Args:
            space_key: Confluence space key (e.g., "TECH")

        Returns:
            List of document chunks
        """
        chunks = []
        try:
            # Import here to avoid dependency if not used
            from requests.auth import HTTPBasicAuth
            import requests

            # Get pages from space
            url = f"{self.base_url}/rest/api/content?spaceKey={space_key}&limit=100&expand=body.storage"
            auth = HTTPBasicAuth(self.username, self.api_token)

            response = requests.get(url, auth=auth, timeout=10)
            response.raise_for_status()

            pages = response.json().get("results", [])

            for page in pages:
                title = page.get("title", "Unknown")
                content = page.get("body", {}).get("storage", {}).get("value", "")
                page_id = page.get("id")

                if content.strip():
                    chunks.append({
                        "content": content,
                        "source": f"Confluence: {title}",
                        "page_id": page_id,
                        "type": "confluence",
                        "loaded_at": datetime.now().isoformat(),
                        "sensitivity": "HIGH"  # Confluence docs often sensitive
                    })

            print(f"✓ Loaded {len(chunks)} pages from Confluence")

        except Exception as e:
            print(f"Error loading Confluence: {e}")

        return chunks
